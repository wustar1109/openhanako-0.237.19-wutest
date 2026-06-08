#!/usr/bin/env python3
from __future__ import annotations

import argparse
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from office_common import (
    DocumentSkillError,
    json_error,
    json_success,
    load_ops,
    prepare_output,
    replace_text_in_containers,
    require_file,
    write_modified_zip,
    write_text_or_json,
    xml_bytes,
)


DOCX_REPLACE_PARTS = (
    "word/document.xml",
    "word/comments.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
)


def is_header_footer_part(name: str) -> bool:
    return (
        name.startswith("word/header") and name.endswith(".xml")
    ) or (
        name.startswith("word/footer") and name.endswith(".xml")
    )


def apply_replace_text(input_path: Path, output_path: Path, op: dict) -> dict:
    find = op.get("find")
    replace = op.get("replace")
    if not isinstance(find, str) or not isinstance(replace, str):
        raise DocumentSkillError("replace_text requires string find and replace fields.", code="invalid_operation")
    max_count = op.get("count")
    if max_count is not None:
        try:
            max_count = int(max_count)
        except Exception as exc:
            raise DocumentSkillError("replace_text count must be an integer.", code="invalid_operation") from exc
        if max_count < 1:
            raise DocumentSkillError("replace_text count must be positive.", code="invalid_operation")

    replacements: dict[str, bytes] = {}
    total = 0
    with zipfile.ZipFile(input_path) as zf:
        for name in zf.namelist():
            if name not in DOCX_REPLACE_PARTS and not is_header_footer_part(name):
                continue
            try:
                root = ET.fromstring(zf.read(name))
            except Exception:
                continue
            changed = replace_text_in_containers(
                root,
                container_names={"p"},
                find=find,
                replace=replace,
                max_count=max_count if total == 0 else (None if max_count is None else max_count - total),
            )
            if changed:
                replacements[name] = xml_bytes(root)
                total += changed
                if max_count is not None and total >= max_count:
                    break

    if not replacements:
        shutil.copyfile(input_path, output_path)
    else:
        write_modified_zip(input_path, output_path, replacements)
    return {"op": "replace_text", "matches": total}


def require_python_docx():
    try:
        import docx  # type: ignore
        return docx
    except Exception as exc:
        raise DocumentSkillError(
            "This DOCX operation requires the optional dependency python-docx.",
            code="missing_dependency",
            details={"package": "python-docx", "original_error": str(exc)},
        ) from exc


def apply_python_docx_op(input_path: Path, output_path: Path, op: dict) -> dict:
    docx = require_python_docx()
    document = docx.Document(str(input_path))
    kind = op["op"]

    if kind == "append_paragraph":
        text = op.get("text")
        if not isinstance(text, str):
            raise DocumentSkillError("append_paragraph requires a text string.", code="invalid_operation")
        paragraph = document.add_paragraph(text)
        style = op.get("style")
        if isinstance(style, str) and style:
            paragraph.style = style
        result = {"op": kind, "paragraphs_added": 1}
    elif kind == "add_table":
        rows = op.get("rows")
        if not isinstance(rows, list) or not rows or not all(isinstance(row, list) for row in rows):
            raise DocumentSkillError("add_table requires rows as a non-empty list of lists.", code="invalid_operation")
        width = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=width)
        for r_idx, row in enumerate(rows):
            for c_idx in range(width):
                table.cell(r_idx, c_idx).text = "" if c_idx >= len(row) else str(row[c_idx])
        style = op.get("style")
        if isinstance(style, str) and style:
            table.style = style
        result = {"op": kind, "rows": len(rows), "cols": width}
    else:
        raise DocumentSkillError(f"Unsupported DOCX operation: {kind}", code="unsupported_operation")

    document.save(str(output_path))
    return result


def run_ops(input_path: Path, output_path: Path, ops: list[dict]) -> list[dict]:
    tmp_dir = Path(tempfile.mkdtemp(prefix="edit-docx-"))
    results: list[dict] = []
    try:
        current = tmp_dir / "step-0.docx"
        shutil.copyfile(input_path, current)
        for idx, op in enumerate(ops, 1):
            next_path = tmp_dir / f"step-{idx}.docx"
            kind = op["op"]
            if kind == "replace_text":
                results.append(apply_replace_text(current, next_path, op))
            elif kind in {"append_paragraph", "add_table"}:
                results.append(apply_python_docx_op(current, next_path, op))
            else:
                raise DocumentSkillError(f"Unsupported DOCX operation: {kind}", code="unsupported_operation")
            current = next_path
        shutil.copyfile(current, output_path)
        return results
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="Apply safe, limited edits to DOCX files.")
    parser.add_argument("input")
    parser.add_argument("output")
    parser.add_argument("--ops", required=True, help="JSON operations file.")
    args = parser.parse_args()

    try:
        input_path = require_file(args.input)
        output_path = prepare_output(input_path, args.output)
        if input_path.suffix.lower() != ".docx" or output_path.suffix.lower() != ".docx":
            raise DocumentSkillError("edit_docx.py requires .docx input and output paths.", code="unsupported_format")
        ops = load_ops(args.ops)
        results = run_ops(input_path, output_path, ops)
        write_text_or_json(None, json_success(input=str(input_path), output=str(output_path), operations=results))
        return 0
    except Exception as exc:
        write_text_or_json(None, json_error(exc))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
